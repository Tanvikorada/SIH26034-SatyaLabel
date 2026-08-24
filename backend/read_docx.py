import sys
import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.replace('\n', ' '))
            full_text.append(' | '.join(row_data))
            
    return '\n'.join(full_text)

if __name__ == '__main__':
    file_path = sys.argv[1]
    out_path = sys.argv[2]
    text = read_docx(file_path)
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(text)
